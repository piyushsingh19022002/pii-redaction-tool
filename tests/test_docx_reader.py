import os
import pytest
import docx
from src.docx_reader import extract_segments
from src.models import TextSegment

@pytest.fixture
def temp_docx_file(tmp_path):
    """Programmatically generates a small test DOCX file with varying formats."""
    doc_path = tmp_path / "test_doc.docx"
    doc = docx.Document()
    
    # Add paragraphs: some valid, some empty, some whitespace
    doc.add_paragraph("First paragraph text.")
    doc.add_paragraph("")  # Empty
    doc.add_paragraph("   ")  # Whitespace
    doc.add_paragraph("Second paragraph text.")
    
    # Add a table (2x2)
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Cell 0,0 text"
    table.cell(0, 1).text = ""  # Empty
    table.cell(1, 0).text = "   "  # Whitespace
    table.cell(1, 1).text = "Cell 1,1 text"
    
    doc.save(str(doc_path))
    return str(doc_path)

def test_paragraph_extraction(temp_docx_file):
    """Verifies that non-empty paragraphs are extracted and indices are preserved."""
    segments = extract_segments(temp_docx_file)
    paragraphs = [s for s in segments if s.segment_type == "paragraph"]
    
    # Empty and whitespace paragraphs should be filtered out
    assert len(paragraphs) == 2
    
    assert paragraphs[0].text == "First paragraph text."
    assert paragraphs[0].paragraph_index == 0
    assert paragraphs[0].segment_type == "paragraph"
    
    assert paragraphs[1].text == "Second paragraph text."
    assert paragraphs[1].paragraph_index == 3  # Index 1 and 2 were empty/whitespace
    assert paragraphs[1].segment_type == "paragraph"

def test_table_cell_extraction(temp_docx_file):
    """Verifies that table cells are extracted and location indices are populated."""
    segments = extract_segments(temp_docx_file)
    cells = [s for s in segments if s.segment_type == "table-cell"]
    
    # Empty/whitespace cells must be skipped
    assert len(cells) == 2
    
    assert cells[0].text == "Cell 0,0 text"
    assert cells[0].table_index == 0
    assert cells[0].row_index == 0
    assert cells[0].cell_index == 0
    assert cells[0].segment_type == "table-cell"
    
    assert cells[1].text == "Cell 1,1 text"
    assert cells[1].table_index == 0
    assert cells[1].row_index == 1
    assert cells[1].cell_index == 1
    assert cells[1].segment_type == "table-cell"

def test_ignores_empty_and_whitespace(temp_docx_file):
    """Verifies that all returned segments contain actual content."""
    segments = extract_segments(temp_docx_file)
    
    # Expecting 2 valid paragraphs + 2 valid cells = 4 segments total
    assert len(segments) == 4
    
    for segment in segments:
        assert len(segment.text.strip()) > 0

def test_real_docx_loading():
    """Verifies that the actual target document exists and can be successfully parsed."""
    real_path = os.path.join("input", "Red Herring Prospectus.docx")
    assert os.path.exists(real_path), f"Expected real DOCX file at {real_path}"
    
    segments = extract_segments(real_path)
    assert len(segments) > 0
    
    # Validate segment types are either 'paragraph' or 'table-cell'
    for s in segments:
        assert s.segment_type in ("paragraph", "table-cell")
        assert len(s.text.strip()) > 0
