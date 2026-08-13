import os
import pathlib
import pytest
import docx
from src.models import TextSegment, PIIEntity, PIIType
from src.docx_redactor import redact_docx

@pytest.fixture
def temp_dir(tmp_path):
    return tmp_path

def test_single_pii_in_paragraph(temp_dir):
    """Verifies that a single PII match in a paragraph is replaced correctly."""
    input_file = temp_dir / "input_single.docx"
    output_file = temp_dir / "output_single.docx"

    doc = docx.Document()
    doc.add_paragraph("John Doe joined today.")
    doc.save(str(input_file))

    segment = TextSegment("John Doe joined today.", "paragraph", paragraph_index=0)
    entity = PIIEntity("John Doe", PIIType.PERSON, 0, 8, 0.90, "ner")
    replacements = [(segment, entity, "Jane Smith")]

    redact_docx(str(input_file), str(output_file), replacements)

    # Validate output
    doc_out = docx.Document(str(output_file))
    assert doc_out.paragraphs[0].text == "Jane Smith joined today."

def test_multiple_pii_in_one_paragraph(temp_dir):
    """Verifies that multiple PII matches in a single paragraph are replaced correctly."""
    input_file = temp_dir / "input_multiple.docx"
    output_file = temp_dir / "output_multiple.docx"

    doc = docx.Document()
    doc.add_paragraph("John Doe works at Google.")
    doc.save(str(input_file))

    segment = TextSegment("John Doe works at Google.", "paragraph", paragraph_index=0)
    ent_person = PIIEntity("John Doe", PIIType.PERSON, 0, 8, 0.90, "ner")
    ent_org = PIIEntity("Google", PIIType.ORGANIZATION, 18, 24, 0.90, "ner")

    replacements = [
        (segment, ent_person, "Jane Smith"),
        (segment, ent_org, "Example Technologies")
    ]

    redact_docx(str(input_file), str(output_file), replacements)

    doc_out = docx.Document(str(output_file))
    assert doc_out.paragraphs[0].text == "Jane Smith works at Example Technologies."

def test_pii_in_table_cell(temp_dir):
    """Verifies that PII inside a table cell is redacted correctly."""
    input_file = temp_dir / "input_table.docx"
    output_file = temp_dir / "output_table.docx"

    doc = docx.Document()
    table = doc.add_table(rows=1, cols=1)
    table.rows[0].cells[0].paragraphs[0].text = "DOB: 01/02/1995"
    doc.save(str(input_file))

    segment = TextSegment("DOB: 01/02/1995", "table-cell", table_index=0, row_index=0, cell_index=0)
    entity = PIIEntity("01/02/1995", PIIType.DOB, 5, 15, 0.90, "regex")

    replacements = [(segment, entity, "15/06/1990")]

    redact_docx(str(input_file), str(output_file), replacements)

    doc_out = docx.Document(str(output_file))
    cell_text = doc_out.tables[0].rows[0].cells[0].paragraphs[0].text
    assert cell_text == "DOB: 15/06/1990"

def test_cross_run_pii_replacement(temp_dir):
    """MANDATORY: Verifies PII split across runs is replaced correctly, keeping context."""
    input_file = temp_dir / "input_cross_run.docx"
    output_file = temp_dir / "output_cross_run.docx"

    doc = docx.Document()
    p = doc.add_paragraph()
    # Add runs manually to force split
    r1 = p.add_run("Contact john@")
    r2 = p.add_run("example")
    r3 = p.add_run(".com today.")
    doc.save(str(input_file))

    # PII spans "john@example.com" -> offsets 8 to 24 in raw text: "Contact john@example.com today."
    segment = TextSegment("Contact john@example.com today.", "paragraph", paragraph_index=0)
    entity = PIIEntity("john@example.com", PIIType.EMAIL, 8, 24, 0.95, "regex")
    replacements = [(segment, entity, "john.doe@example.com")]

    redact_docx(str(input_file), str(output_file), replacements)

    doc_out = docx.Document(str(output_file))
    assert doc_out.paragraphs[0].text == "Contact john.doe@example.com today."

def test_formatting_preservation(temp_dir):
    """Verifies that text formatting (e.g. bold) is preserved in unaffected portions of runs."""
    input_file = temp_dir / "input_format.docx"
    output_file = temp_dir / "output_format.docx"

    doc = docx.Document()
    p = doc.add_paragraph()
    r1 = p.add_run("Yesterday, ")
    r2 = p.add_run("John Doe")
    r2.bold = True
    r3 = p.add_run(" joined.")
    doc.save(str(input_file))

    segment = TextSegment("Yesterday, John Doe joined.", "paragraph", paragraph_index=0)
    entity = PIIEntity("John Doe", PIIType.PERSON, 11, 19, 0.90, "ner")
    replacements = [(segment, entity, "Jane Smith")]

    redact_docx(str(input_file), str(output_file), replacements)

    doc_out = docx.Document(str(output_file))
    p_out = doc_out.paragraphs[0]
    
    # Assert formatting remains intact
    assert len(p_out.runs) == 3
    assert p_out.runs[0].text == "Yesterday, "
    assert p_out.runs[1].text == "Jane Smith"
    assert p_out.runs[1].bold is True
    assert p_out.runs[2].text == " joined."

def test_empty_replacement_list(temp_dir):
    """Verifies that an empty replacements list produces identical document text."""
    input_file = temp_dir / "input_empty_rep.docx"
    output_file = temp_dir / "output_empty_rep.docx"

    doc = docx.Document()
    doc.add_paragraph("This is some text.")
    doc.save(str(input_file))

    redact_docx(str(input_file), str(output_file), [])

    doc_out = docx.Document(str(output_file))
    assert doc_out.paragraphs[0].text == "This is some text."

def test_original_document_remains_unchanged(temp_dir):
    """Verifies that redacting a copy does not affect the original source file."""
    input_file = temp_dir / "input_unchanged.docx"
    output_file = temp_dir / "output_redacted.docx"

    doc = docx.Document()
    doc.add_paragraph("Original text.")
    doc.save(str(input_file))

    # Read modify timestamp
    orig_stat = os.stat(str(input_file))

    segment = TextSegment("Original text.", "paragraph", paragraph_index=0)
    entity = PIIEntity("Original", PIIType.ORGANIZATION, 0, 8, 0.50, "ner")
    replacements = [(segment, entity, "Redacted")]

    redact_docx(str(input_file), str(output_file), replacements)

    # Check original document text is still unchanged
    doc_orig = docx.Document(str(input_file))
    assert doc_orig.paragraphs[0].text == "Original text."
