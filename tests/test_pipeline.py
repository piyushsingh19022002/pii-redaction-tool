import os
import pytest
import docx
from src.models import PIIType
from src.pipeline import PIIRedactionPipeline
from src.detectors.email import EmailDetector
from src.detectors.phone import PhoneDetector

@pytest.fixture
def temp_dir(tmp_path):
    return tmp_path

def test_pipeline_end_to_end_paragraph_pii(temp_dir):
    """Verifies end-to-end PII detection and redaction in paragraphs using injected detectors."""
    input_file = temp_dir / "input_pipeline.docx"
    output_file = temp_dir / "output_pipeline.docx"

    doc = docx.Document()
    doc.add_paragraph("Contact John at john@example.com.")
    doc.save(str(input_file))

    # Inject EmailDetector only to avoid spaCy initialization overhead in this unit test
    pipeline = PIIRedactionPipeline(detectors=[EmailDetector()])
    result = pipeline.run(str(input_file), str(output_file))

    assert result.segments_processed == 1
    assert result.candidates_detected == 1
    assert result.candidates_accepted == 1
    assert result.counts_by_type["EMAIL"] == 1

    # Verify output file content
    doc_out = docx.Document(str(output_file))
    assert "@example.com" in doc_out.paragraphs[0].text
    assert "john@example.com" not in doc_out.paragraphs[0].text

def test_pipeline_end_to_end_table_pii(temp_dir):
    """Verifies end-to-end PII detection and redaction in table cells."""
    input_file = temp_dir / "input_table.docx"
    output_file = temp_dir / "output_table.docx"

    doc = docx.Document()
    table = doc.add_table(rows=1, cols=1)
    table.rows[0].cells[0].paragraphs[0].text = "Reach us at +91 9876543210."
    doc.save(str(input_file))

    pipeline = PIIRedactionPipeline(detectors=[PhoneDetector()])
    result = pipeline.run(str(input_file), str(output_file))

    assert result.segments_processed == 1
    assert result.candidates_detected == 1
    assert result.candidates_accepted == 1
    assert result.counts_by_type["PHONE"] == 1

    doc_out = docx.Document(str(output_file))
    cell_text = doc_out.tables[0].rows[0].cells[0].paragraphs[0].text
    assert "+91" in cell_text
    assert "9876543210" not in cell_text

def test_pipeline_pseudonymizer_shared_instance(temp_dir):
    """Verifies that a single pseudonymizer is shared and maps duplicates consistently across segments."""
    input_file = temp_dir / "input_consistency.docx"
    output_file = temp_dir / "output_consistency.docx"

    doc = docx.Document()
    doc.add_paragraph("Email is john@example.com.")
    doc.add_paragraph("Second email is john@example.com.")
    doc.save(str(input_file))

    pipeline = PIIRedactionPipeline(detectors=[EmailDetector()])
    result = pipeline.run(str(input_file), str(output_file))

    assert result.segments_processed == 2
    assert result.candidates_accepted == 2

    # Verify both occurrences were replaced with the EXACT same fake email
    doc_out = docx.Document(str(output_file))
    email_replacement_1 = doc_out.paragraphs[0].text.split("is ")[1].strip(".")
    email_replacement_2 = doc_out.paragraphs[1].text.split("is ")[1].strip(".")
    
    assert email_replacement_1 == email_replacement_2
    assert email_replacement_1.endswith("@example.com")

def test_pipeline_missing_input_file(temp_dir):
    """Verifies that running the pipeline with a non-existent input path raises FileNotFoundError."""
    pipeline = PIIRedactionPipeline(detectors=[EmailDetector()])
    with pytest.raises(FileNotFoundError):
        pipeline.run("non_existent_file.docx", str(temp_dir / "output.docx"))

def test_pipeline_original_file_remains_unchanged(temp_dir):
    """Verifies that the original file is not modified or mutated during pipeline execution."""
    input_file = temp_dir / "input_unchanged.docx"
    output_file = temp_dir / "output_redacted.docx"

    doc = docx.Document()
    doc.add_paragraph("Original email is john@example.com.")
    doc.save(str(input_file))

    pipeline = PIIRedactionPipeline(detectors=[EmailDetector()])
    pipeline.run(str(input_file), str(output_file))

    # Read original document and verify content is untouched
    doc_orig = docx.Document(str(input_file))
    assert doc_orig.paragraphs[0].text == "Original email is john@example.com."
