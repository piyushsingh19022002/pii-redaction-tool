import os
import re
import pathlib
import docx
from typing import List, Tuple, Dict
from src.models import TextSegment, PIIEntity

def find_raw_offsets(
    raw_text: str,
    normalized_text: str,
    cand_start: int,
    cand_end: int,
    cand_text: str
) -> Tuple[int, int]:
    """Maps the start/end offsets from normalized_text back to the raw_text.

    Uses occurrence indexing of the target substring to handle any whitespace or
    margin character discrepancies caused by normalization.

    Args:
        raw_text: The original text of the paragraph or cell.
        normalized_text: The normalized text representation.
        cand_start: Start character index in normalized_text.
        cand_end: End character index in normalized_text.
        cand_text: The matched PII string.

    Returns:
        A tuple of (start, end) offsets in the raw_text.
    """
    if not raw_text or not cand_text:
        return cand_start, cand_end

    # 1. Count how many times cand_text appears in normalized_text before cand_start
    occurrence_index = 0
    pos = normalized_text.find(cand_text)
    while pos != -1 and pos < cand_start:
        occurrence_index += 1
        pos = normalized_text.find(cand_text, pos + 1)

    # 2. Find the same occurrence index of cand_text in raw_text
    raw_pos = -1
    for _ in range(occurrence_index + 1):
        raw_pos = raw_text.find(cand_text, raw_pos + 1)
        if raw_pos == -1:
            break

    if raw_pos != -1:
        return raw_pos, raw_pos + len(cand_text)
    else:
        # Fallback if raw lookup fails
        return cand_start, cand_end

def redact_paragraph(
    paragraph: docx.text.paragraph.Paragraph,
    entity_replacements: List[Tuple[PIIEntity, str]],
    segment_normalized_text: str
) -> None:
    """Replaces PII spans across runs in a paragraph using right-to-left replacement.

    Args:
        paragraph: The python-docx Paragraph element.
        entity_replacements: List of tuples containing (PIIEntity, replacement_string).
        segment_normalized_text: Normalized text corresponding to the paragraph.
    """
    if not paragraph.runs or not entity_replacements:
        return

    # 1. Build visible text of paragraph from runs
    raw_text = "".join(run.text for run in paragraph.runs)

    # 2. Map normalized offsets of each entity to raw_text offsets
    mapped_replacements = []
    for entity, replacement in entity_replacements:
        raw_start, raw_end = find_raw_offsets(
            raw_text=raw_text,
            normalized_text=segment_normalized_text,
            cand_start=entity.start,
            cand_end=entity.end,
            cand_text=entity.text
        )
        mapped_replacements.append((raw_start, raw_end, replacement))

    # 3. Sort replacements from right to left (descending by raw_start)
    mapped_replacements.sort(key=lambda x: x[0], reverse=True)

    # 4. Map run boundaries
    run_ranges = []
    current_idx = 0
    for run in paragraph.runs:
        run_len = len(run.text)
        run_ranges.append((current_idx, current_idx + run_len, run))
        current_idx += run_len

    # 5. Apply replacements from right to left
    for start, end, replacement in mapped_replacements:
        # Find runs affected by [start, end)
        affected_runs = []
        for run_start, run_end, run in run_ranges:
            # Overlap check
            if run_start < end and start < run_end:
                affected_runs.append((run_start, run_end, run))

        if not affected_runs:
            continue

        # First affected run gets the replacement text
        for i, (run_start, run_end, run) in enumerate(affected_runs):
            run_text = run.text
            rel_start = max(0, start - run_start)
            rel_end = max(0, end - run_start)

            if i == 0:
                run.text = run_text[:rel_start] + replacement + run_text[rel_end:]
            else:
                run.text = run_text[:rel_start] + "" + run_text[rel_end:]

        # Rebuild run boundaries to keep offsets accurate for the next replacement
        run_ranges = []
        current_idx = 0
        for run in paragraph.runs:
            run_len = len(run.text)
            run_ranges.append((current_idx, current_idx + run_len, run))
            current_idx += run_len

def redact_docx(
    input_path: str,
    output_path: str,
    replacements: List[Tuple[TextSegment, PIIEntity, str]]
) -> None:
    """Redacts PII entities in a DOCX document and saves the result to a new file.

    Supports paragraphs and table cells, handling cross-run PII boundaries and
    preserving formatting.

    Args:
        input_path: Path to the input DOCX file.
        output_path: Path where the redacted DOCX file will be saved.
        replacements: List of Tuples containing (TextSegment, PIIEntity, replacement_string).
    """
    # Open document
    doc = docx.Document(input_path)

    # Group replacements by segment type and coordinates
    paragraph_replacements: Dict[int, List[Tuple[PIIEntity, str, str]]] = {}
    table_cell_replacements: Dict[Tuple[int, int, int], List[Tuple[PIIEntity, str, str]]] = {}

    for segment, entity, replacement in replacements:
        if segment.segment_type == "paragraph":
            idx = segment.paragraph_index
            if idx is not None:
                if idx not in paragraph_replacements:
                    paragraph_replacements[idx] = []
                paragraph_replacements[idx].append((entity, replacement, segment.normalized_text or ""))
        elif segment.segment_type == "table-cell":
            t_idx = segment.table_index
            r_idx = segment.row_index
            c_idx = segment.cell_index
            if t_idx is not None and r_idx is not None and c_idx is not None:
                cell_key = (t_idx, r_idx, c_idx)
                if cell_key not in table_cell_replacements:
                    table_cell_replacements[cell_key] = []
                table_cell_replacements[cell_key].append((entity, replacement, segment.normalized_text or ""))

    # 1. Redact normal paragraphs
    for idx, entity_replacements in paragraph_replacements.items():
        if idx < len(doc.paragraphs):
            # Convert tuples
            reps_list = [(entity, rep) for entity, rep, _ in entity_replacements]
            norm_text = entity_replacements[0][2]
            redact_paragraph(doc.paragraphs[idx], reps_list, norm_text)

    # 2. Redact table cell paragraphs
    for cell_key, entity_replacements in table_cell_replacements.items():
        t_idx, r_idx, c_idx = cell_key
        if t_idx < len(doc.tables):
            table = doc.tables[t_idx]
            if r_idx < len(table.rows):
                row = table.rows[r_idx]
                if c_idx < len(row.cells):
                    cell = row.cells[c_idx]
                    segment_normalized_text = entity_replacements[0][2]

                    # Map cell-level replacements to paragraph relative offsets
                    paragraphs = cell.paragraphs
                    if not paragraphs:
                        continue

                    # Construct cell raw text and map boundaries
                    cell_raw_text = ""
                    para_ranges = []
                    for p in paragraphs:
                        p_len = len(p.text)
                        start_idx = len(cell_raw_text)
                        if cell_raw_text:
                            cell_raw_text += "\n"
                            start_idx += 1
                        cell_raw_text += p.text
                        para_ranges.append((start_idx, start_idx + p_len, p))

                    para_replacements: Dict[docx.text.paragraph.Paragraph, List[Tuple[PIIEntity, str]]] = {}
                    for entity, replacement, _ in entity_replacements:
                        raw_start, raw_end = find_raw_offsets(
                            raw_text=cell_raw_text,
                            normalized_text=segment_normalized_text,
                            cand_start=entity.start,
                            cand_end=entity.end,
                            cand_text=entity.text
                        )

                        # Find containing paragraph
                        target_para = None
                        rel_start = 0
                        rel_end = 0
                        for p_start, p_end, p in para_ranges:
                            if p_start <= raw_start < p_end:
                                target_para = p
                                rel_start = raw_start - p_start
                                rel_end = raw_end - p_start
                                break

                        if target_para is None:
                            target_para = paragraphs[0]
                            rel_start = 0
                            rel_end = len(target_para.text)

                        rel_entity = PIIEntity(
                            text=entity.text,
                            entity_type=entity.entity_type,
                            start=rel_start,
                            end=rel_end,
                            confidence=entity.confidence,
                            source=entity.source
                        )

                        if target_para not in para_replacements:
                            para_replacements[target_para] = []
                        para_replacements[target_para].append((rel_entity, replacement))

                    # Apply redaction to each paragraph inside the cell
                    for p, p_reps in para_replacements.items():
                        redact_paragraph(p, p_reps, p.text)

    # Save redacted file
    out_path = pathlib.Path(output_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
