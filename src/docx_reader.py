import docx
from typing import List
from src.models import TextSegment
from src.normalizer import normalize_text

def extract_segments(file_path: str) -> List[TextSegment]:
    """Reads a DOCX document and extracts its non-empty paragraphs and table cells.

    Preserves the document index of paragraphs and row/cell coordinates
    for tables to allow correct identification and mapping in future steps.

    Args:
        file_path: Path to the input DOCX file.

    Returns:
        A list of TextSegment dataclasses containing the extracted texts and metadata.
    """
    doc = docx.Document(file_path)
    segments: List[TextSegment] = []

    # Extract paragraphs
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text:  # Ignore empty or whitespace-only paragraphs
            segments.append(TextSegment(
                text=text,
                segment_type="paragraph",
                paragraph_index=idx,
                normalized_text=normalize_text(text)
            ))

    # Extract tables
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                if text:  # Ignore empty or whitespace-only table cells
                    segments.append(TextSegment(
                        text=text,
                        segment_type="table-cell",
                        table_index=table_idx,
                        row_index=row_idx,
                        cell_index=cell_idx,
                        normalized_text=normalize_text(text)
                    ))

    return segments

if __name__ == "__main__":
    import os

    input_path = os.path.join("input", "Red Herring Prospectus.docx")
    
    if not os.path.exists(input_path):
        print(f"Error: '{input_path}' not found. Please ensure the file is placed in the input/ folder.")
    else:
        print(f"Extracting structure from: {input_path}")
        try:
            segments = extract_segments(input_path)
            
            # Group segment types for statistics
            paragraphs = [s for s in segments if s.segment_type == "paragraph"]
            cells = [s for s in segments if s.segment_type == "table-cell"]
            table_count = len(set(s.table_index for s in cells if s.table_index is not None))
            
            print("\n==================================================")
            print("DOCX EXTRACTION SUMMARY")
            print("==================================================")
            print(f"Total Extracted Segments: {len(segments)}")
            print(f"  - Paragraph Segments:   {len(paragraphs)}")
            print(f"  - Table-cell Segments:  {len(cells)}")
            print(f"  - Total Unique Tables:  {table_count}")
            print("==================================================\n")
            
            # Print paragraph sample with raw vs normalized text
            print("SAMPLE PARAGRAPH SEGMENTS (RAW vs NORMALIZED):")
            print("--------------------------------------------------")
            for s in paragraphs[:2]:
                print(f"[Para Index {s.paragraph_index}]")
                print(f"  RAW:        {repr(s.text[:100])}")
                print(f"  NORMALIZED: {repr(s.normalized_text[:100])}")
                print()
            print("--------------------------------------------------\n")
            
            # Print table cell sample with raw vs normalized text
            print("SAMPLE TABLE-CELL SEGMENTS (RAW vs NORMALIZED):")
            print("--------------------------------------------------")
            for s in cells[:2]:
                print(f"[Table {s.table_index}, Row {s.row_index}, Cell {s.cell_index}]")
                print(f"  RAW:        {repr(s.text[:100])}")
                print(f"  NORMALIZED: {repr(s.normalized_text[:100])}")
                print()
            print("--------------------------------------------------")
            
        except Exception as e:
            print(f"Failed to process document: {e}")
